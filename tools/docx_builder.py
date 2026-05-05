"""
Reusable resume + cover letter builder.

Usage (from a job script):
    from tools.docx_builder import ResumeBuilder, CoverLetterBuilder

    r = ResumeBuilder(accent=(31,78,121))
    r.header()
    r.summary("...")
    r.skills([("Languages", "Java, Python"), ...])
    r.experience([
        {
            "company": "...", "title": "...", "dates": "...", "location": "...",
            "bullets": ["...", "..."]
        },
        ...
    ])
    r.education()
    r.save("path/to/Name_Company.docx")

    cl = CoverLetterBuilder(accent=(31,78,121))
    cl.header()
    cl.salutation(date="March 17, 2026", to="Hiring Team", company="...", city="...", role="...")
    cl.body([para1, para2, ...])
    cl.sign_off()
    cl.save("path/to/message_Company.docx")
"""

import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Load user config from personal-docs/user.json ────────────────────────────
_root     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_cfg_path = os.path.join(_root, "personal-docs", "user.json")

try:
    with open(_cfg_path) as _f:
        _cfg = json.load(_f)

    USER_NAME      = _cfg["name"]
    CONTACT_LINE   = _cfg["contact_line"]
    WORK_AUTH_PARA = _cfg.get("work_auth_para", "")
    LINKEDIN_PARA  = _cfg.get("linkedin_para", "")
    EDUCATION      = [
        (e["degree"], e["institution"], e["location"], e["year"])
        for e in _cfg["education"]
    ]
    _COVER_CONTACT = "  |  ".join(filter(None, [
        _cfg.get("phone", ""),
        _cfg.get("email", ""),
        _cfg.get("linkedin", ""),
        _cfg.get("portfolio", ""),
        _cfg.get("location", ""),
    ]))

except FileNotFoundError:
    raise RuntimeError(
        "personal-docs/user.json not found. Run /setup to initialize your profile."
    )


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _sf(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _hr(doc, color_hex, sb=1, sa=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _rgb_to_hex(rgb):
    return "{:02X}{:02X}{:02X}".format(*rgb)


# ── ResumeBuilder ─────────────────────────────────────────────────────────────

class ResumeBuilder:
    def __init__(self, accent=(68, 114, 196)):
        self.accent = accent
        self.hex    = _rgb_to_hex(accent)
        self.doc    = Document()
        for s in self.doc.sections:
            s.top_margin    = Inches(0.45)
            s.bottom_margin = Inches(0.45)
            s.left_margin   = Inches(0.6)
            s.right_margin  = Inches(0.6)

    def _sec(self, title):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(title.upper())
        _sf(r, size=10, bold=True, color=self.accent)
        _hr(self.doc, self.hex)

    def _bullet(self, text, size=9.5, sa=1):
        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(sa)
        p.paragraph_format.left_indent       = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(u"•  " + text)
        _sf(r, size=size)

    def header(self):
        np = self.doc.add_paragraph()
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(0)
        np.paragraph_format.space_after  = Pt(2)
        r = np.add_run(USER_NAME.upper())
        _sf(r, size=18, bold=True, color=self.accent)

        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(2)
        r = cp.add_run(CONTACT_LINE)
        _sf(r, size=9, color=(89, 89, 89))
        _hr(self.doc, self.hex)

    def summary(self, text):
        self._sec("Professional Summary")
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        _sf(r, size=9.5)

    def skills(self, rows):
        """rows: list of (label, value) tuples"""
        self._sec("Technical Skills")
        for label, val in rows:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            r1 = p.add_run(label + ":  ")
            _sf(r1, size=9.5, bold=True)
            r2 = p.add_run(val)
            _sf(r2, size=9.5)

    def experience(self, roles):
        """
        roles: list of dicts with keys:
            company, title, dates, location, bullets (list of str)
        """
        self._sec("Professional Experience")
        for role in roles:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_ALIGN_PARAGRAPH.RIGHT)
            r1 = p.add_run(role["company"]); _sf(r1, size=10, bold=True)
            r2 = p.add_run(" | " + role["title"]); _sf(r2, size=9.5)
            r3 = p.add_run("\t" + role["dates"]); _sf(r3, size=9, color=(89,89,89))

            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(1)
            r4 = p2.add_run(role["location"]); _sf(r4, size=9, color=(89,89,89))

            for b in role["bullets"]:
                self._bullet(b)

    def education(self, rows=None):
        self._sec("Education")
        for degree, inst, loc, yr in (rows or EDUCATION):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(1)
            r1 = p.add_run(degree + " — "); _sf(r1, size=9.5, bold=True)
            r2 = p.add_run(inst + ", " + loc + " | " + yr); _sf(r2, size=9.5)

    def save(self, path):
        self.doc.save(path)
        print("Resume saved:", path)


# ── CoverLetterBuilder ────────────────────────────────────────────────────────

class CoverLetterBuilder:
    def __init__(self, accent=(68, 114, 196)):
        self.accent = accent
        self.hex    = _rgb_to_hex(accent)
        self.doc    = Document()
        for s in self.doc.sections:
            s.top_margin    = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin   = Inches(1.0)
            s.right_margin  = Inches(1.0)

    def _p(self, text, size=11, bold=False, sa=8, color=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(sa)
        if text:
            r = p.add_run(text)
            _sf(r, size=size, bold=bold, color=color)
        return p

    def _hr(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), self.hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    def header(self):
        np = self.doc.add_paragraph()
        np.paragraph_format.space_before = Pt(0)
        np.paragraph_format.space_after  = Pt(2)
        r = np.add_run(USER_NAME.upper())
        _sf(r, size=16, bold=True, color=self.accent)

        cp = self.doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(4)
        r = cp.add_run(_COVER_CONTACT)
        _sf(r, size=9.5, color=(89,89,89))
        self._hr()

    def salutation(self, date, to, company, city, role):
        self._p(date, sa=6)
        self._p(to, sa=2)
        self._p(company, sa=2)
        self._p(city, sa=12)
        self._p("Re: " + role, bold=True, sa=12)

    def body(self, paragraphs):
        """paragraphs: list of strings, each becomes one paragraph"""
        for text in paragraphs:
            self._p(text, sa=10)

    def sign_off(self, closing="Resume is attached. Available for a call at your convenience."):
        self._p(closing, sa=20)
        self._p(USER_NAME, bold=True, sa=2)
        self._p(_cfg.get("phone", ""), sa=2)
        self._p(_cfg.get("email", ""), sa=2)
        if _cfg.get("linkedin"):
            self._p(_cfg["linkedin"], sa=2)
        if _cfg.get("portfolio"):
            self._p(_cfg["portfolio"], sa=6)

    def save(self, path):
        self.doc.save(path)
        print("Cover letter saved:", path)
